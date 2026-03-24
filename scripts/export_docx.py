#!/usr/bin/env python3
"""
将 Markdown 标书文件导出为 Word (.docx) 格式

用法：
    python export_docx.py /path/to/标书.md
    python export_docx.py /path/to/标书.md -o output.docx
"""

import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_styles(doc):
    """配置文档样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # 标题样式
    heading_configs = [
        ('Heading 1', '黑体', 22, True),
        ('Heading 2', '黑体', 16, True),
        ('Heading 3', '黑体', 14, True),
        ('Heading 4', '宋体', 12, True),
    ]
    for style_name, font_name, size, bold in heading_configs:
        s = doc.styles[style_name]
        s.font.name = font_name
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, header_row, data_rows):
    """添加表格"""
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = text.strip()
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, '2563EB')

    # 数据行
    for row_data in data_rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            if i < cols:
                cell = row.cells[i]
                cell.text = text.strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 斑马纹
    for idx, row in enumerate(table.rows[1:], 1):
        if idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, 'F1F5F9')

    doc.add_paragraph()  # 表格后空行
    return table


def add_formatted_paragraph(doc, text, style=None):
    """添加带格式的段落（支持粗体、行内代码）"""
    p = doc.add_paragraph(style=style)

    # 解析行内格式
    parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 30, 60)
        else:
            run = p.add_run(part)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return p


def md_to_docx(md_text, output_path):
    """将 Markdown 转换为 Word 文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    setup_styles(doc)

    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_content = []
    in_list = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_content = []
                i += 1
                continue
            else:
                # 输出代码块
                if code_content:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                in_code = False
                code_content = []
                i += 1
                continue

        if in_code:
            code_content.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            in_list = False
            i += 1
            continue

        # 水平线
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # 表格
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                header = [c.strip() for c in table_lines[0].strip('|').split('|')]
                data = []
                for tl in table_lines[2:]:  # 跳过分隔行
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    data.append(cells)
                add_table(doc, header, data)
            continue

        # 标题
        heading = re.match(r'^(#{1,4})\s+(.+)', line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            # 清理粗体标记
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_heading(text, level=level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 引用块（跳过 GitHub alerts 标记）
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                content = lines[i].strip()[1:].strip()
                if not re.match(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', content):
                    if content:
                        quote_lines.append(content)
                i += 1
            if quote_lines:
                for ql in quote_lines:
                    p = add_formatted_paragraph(doc, ql)
                    p.paragraph_format.left_indent = Cm(1)
                    pf = p.paragraph_format
                    pf.space_before = Pt(3)
                    pf.space_after = Pt(3)
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*]\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            p = add_formatted_paragraph(doc, text, style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Cm(1.5)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if ol_match:
            text = ol_match.group(2)
            p = add_formatted_paragraph(doc, text, style='List Number')
            i += 1
            continue

        # 普通段落
        text = line.strip()
        if text:
            add_formatted_paragraph(doc, text)
        i += 1

    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Markdown 标书 → Word 导出")
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")

    # 去掉 YAML frontmatter
    md_text = re.sub(r'^---\n.*?\n---\n', '', md_text, flags=re.DOTALL)

    output_path = Path(args.output) if args.output else input_path.with_suffix('.docx')
    md_to_docx(md_text, output_path)
    print(f"✅ 已导出: {output_path}")


if __name__ == "__main__":
    main()
